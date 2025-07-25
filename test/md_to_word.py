import markdown
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def md_to_word(md_file_path, docx_file_path=None):
    """
    将Markdown文件转换为Word文档

    参数:
        md_file_path: Markdown文件路径（.txt格式）
        docx_file_path: 输出Word文档路径，默认为与输入文件同名的.docx文件
    """
    # 如果未指定输出路径，则使用输入文件名并替换扩展名为.docx
    if not docx_file_path:
        docx_file_path = md_file_path.rsplit('.', 1)[0] + '.docx'

    # 读取Markdown文件内容
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 创建Word文档
    doc = Document()

    # 配置字体样式
    # 正文样式
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = '宋体'
    normal_font.size = Pt(11)
    # 设置中文字体
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题样式 - 根据Markdown标题级别设置不同大小
    title_sizes = [Pt(20), Pt(17), Pt(15), Pt(13), Pt(11), Pt(11)]  # 对应h1到h6

    for level in range(1, 7):
        heading_style = doc.styles[f'Heading {level}']
        heading_font = heading_style.font
        heading_font.name = '宋体'
        heading_font.size = title_sizes[level-1]
        heading_font.bold = False
        # 设置中文字体
        heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 标题左对齐
        heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 将Markdown转换为HTML，然后处理内容
    html = markdown.markdown(md_content)

    # 简单解析HTML并添加到Word文档
    # 这里使用简单的字符串处理，更复杂的场景可以考虑使用BeautifulSoup
    lines = html.split('\n')
    current_paragraph = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 处理标题
        heading_level = None
        for i in range(1, 7):
            if line.startswith(f'<h{i}>') and line.endswith(f'</h{i}>'):
                heading_level = i
                break

        if heading_level:
            text = line[3+heading_level-1 : -(4+heading_level-1)]  # 提取标题文本
            doc.add_heading(text, level=heading_level)
            continue

        # 处理段落
        if line.startswith('<p>') and line.endswith('</p>'):
            text = line[3:-4]
            current_paragraph = doc.add_paragraph(text)
            current_paragraph.style = 'Normal'
            continue

        # 处理无序列表
        if line.startswith('<li>') and line.endswith('</li>'):
            text = line[4:-5]
            if current_paragraph is None or current_paragraph.style.name != 'List Bullet':
                current_paragraph = doc.add_paragraph(text, style='List Bullet')
            else:
                current_paragraph.add_run('\n' + text)
            continue

    # 保存Word文档
    doc.save(docx_file_path)
    print(f"转换完成，文件已保存至: {docx_file_path}")

if __name__ == "__main__":

    md_to_word("context.txt", "example.docx")

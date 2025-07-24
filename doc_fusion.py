import os
import argparse
import json
import requests
from docx import Document
from tqdm import tqdm
from typing import List, Tuple
import html2docx


BASE_URL = "http://192.168.1.13:8000/v1"
API_KEY = "2JysWWdHfyKvp2AsGYznw7pwPfkwDehtPZHEtj26GIA"
MODEL_NAME = "qwen-max"

# 请求头
headers = {
    "Content-Type": "application/json",
    "Authorization": f"Bearer {API_KEY}"
}

def read_table(table) -> str:
    """读取表格内容并转换为文本格式"""
    table_text = []
    # 记录表头和内容分隔符
    table_text.append("=== 表格开始 ===")

    # 遍历表格的每一行
    for row_idx, row in enumerate(table.rows):
        row_data = []
        # 遍历行中的每个单元格
        for cell in row.cells:
            # 提取单元格文本，处理可能的换行
            cell_text = "\n".join([p.text for p in cell.paragraphs]).strip()
            row_data.append(cell_text)

        # 用|分隔单元格内容，便于识别
        table_text.append("|".join(row_data))

        # 在表头下方添加分隔线（如果有表头）
        if row_idx == 0 and len(table.rows) > 1:
            table_text.append("|".join(["---"] * len(row_data)))

    table_text.append("=== 表格结束 ===")
    return "\n".join(table_text)

def read_docx(file_path: str) -> str:
    """读取docx文档内容，包括段落和表格"""
    # 创建Document对象读取Word文档
    doc = Document(file_path)
    full_text = []

    # 遍历文档中的每个元素，区分段落和表格
    for element in doc.element.body:
        # 检查元素是否为段落
        if element.tag.endswith('p'):
            # 找到对应的段落对象
            para = next((p for p in doc.paragraphs if p._element == element), None)
            if para and para.text.strip():
                full_text.append(para.text)
        # 检查元素是否为表格
        elif element.tag.endswith('tbl'):
            # 找到对应的表格对象
            table = next((t for t in doc.tables if t._element == element), None)
            if table:
                # 读取表格内容并添加到文本中
                table_content = read_table(table)
                full_text.append(table_content)

    # 将所有内容用换行符连接成一个字符串并返回
    return '\n'.join(full_text)

def create_prompt(files_content: List[Tuple[str, str]], instruction: str = "") -> str:
    """创建提示词"""
    # 初始化提示词，介绍后续是多个文档的内容
    prompt = "以下是多个文档的内容，包含段落和表格（表格用特殊标记标识）:\n\n"

    # 遍历每个文档，将文档名和内容添加到提示词中
    for i, (filename, content) in enumerate(files_content):
        prompt += f"文档 {i+1} ({filename}):\n{content}\n\n"

    # 如果提供了额外说明，将其添加到提示词中
    if instruction:
        prompt += f"额外说明: {instruction}\n\n"

    # 要求模型根据以下格式进行融合重写
    prompt += ("输出格式要求：标题：使用方正小标宋_GBK，字号为二号，加粗。_GBK，字号为二号，加粗。一级标题：格式为：一、XXX，方正黑体_GBK，三号，首行缩进值2"
               "字符，左对齐，行间距28磅，与正文之间不空行，不加粗。二级标题：格式为：（二）XXX，楷体，三号，首行缩进值2字符，左对齐，行间距28"
               "磅，与正文之间不空行，不加粗。三级标题：格式为：1.XXX，仿宋，三号，首行缩进值2字符，行间距28磅，左对齐，与正文之间不空行，不加粗。四级标题：格式为：（1"
               "）XXX，仿宋，三号，首行缩进值2字符，左对齐，行间距28磅，与正文之间不空行，不加粗。正文格式：仿宋三号，首行缩进值2字符，左对齐，行间距28磅。请按照我给的格式进行排版，输出 html 格式。")
    return prompt

def call_llm(prompt: str) -> str:
    """调用LLM API生成内容"""

    url = f"{BASE_URL}/chat/completions"
    payload = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system", "content": "你是一个专业的文档融合专家。请根据用户提供的多个文档内容，进行融合重写，确保内容连贯、逻辑清晰，并且保留原文的核心信息。对于表格数据，要准确保留数据关系和关键信息。"},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3,
        "max_tokens": 4000
    }

    try:
        print(f"发送请求到API，提示词长度: {len(prompt)} 字符")
        response = requests.post(url, headers=headers, data=json.dumps(payload))
        response.raise_for_status()
        result = response.json()["choices"][0]["message"]["content"].strip()

        # 调试输出
        print(f"API返回内容长度: {len(result)} 字符")
        if len(result) < 10:  # 如果内容非常短，可能有问题
            print(f"警告: API返回的内容异常短: {result[:50]}...")

        return result
    except Exception as e:
        print(f"API调用错误: {e}")
        return None

def save_to_html(content: str, output_file: str) -> None:
    """
    将给定的字符串内容保存为HTML文件

    参数:
        content (str): 要保存的HTML内容
        output_file (str): 输出HTML文件的路径

    返回:
        None
    """
    try:
        with open(output_file, 'w', encoding='utf-8') as file:
            file.write(content)
        print(f"内容已成功保存到 {output_file}")
    except IOError as e:
        print(f"保存文件时出错: {e}")
def main():
    # 创建命令行参数解析器，设置工具描述
    parser = argparse.ArgumentParser(description="多文档融合重写工具")
    # 添加输入文件参数 - 允许多个docx文件路径
    parser.add_argument("-i", "--input", nargs="+", help="输入的docx文件路径")
    # 添加输出文件参数 - 指定融合后的文档保存路径
    parser.add_argument("-o", "--output", help="输出文件的路径")
    # 添加额外说明参数 - 可以提供给模型的额外指令
    parser.add_argument("-t", "--instruction", help="额外说明或指令")
    # 解析命令行参数
    args = parser.parse_args()
    # 处理默认输入文件 - 如果未指定输入文件，使用当前目录下的所有docx文件
    if not args.input:
        print("未指定输入文件，将使用当前目录下的所有docx文件")
        current_dir = os.getcwd()
        args.input = [os.path.join(current_dir, f) for f in os.listdir(current_dir) if f.lower().endswith('.docx')]

        # 检查是否找到docx文件
        if not args.input:
            print("错误: 当前目录下没有找到docx文件")
            return

    # 处理默认输出文件 - 如果未指定输出文件，使用默认名称
    if not args.output:
        print("未指定输出文件，将使用默认名称")
        current_dir = os.getcwd()
        args.output = os.path.join(current_dir, "融合文档.html")

    # 检查输入文件是否存在 - 收集不存在的文件
    invalid_files = [f for f in args.input if not os.path.exists(f)]
    if invalid_files:
        print(f"错误: 以下文件不存在: {', '.join(invalid_files)}")
        return

    # 检查输出文件目录是否存在 - 如果不存在则创建
    output_dir = os.path.dirname(args.output)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 读取所有文档内容 - 使用进度条显示读取进度
    print("正在读取文档...")
    files_content = []
    for file_path in tqdm(args.input):
        content = read_docx(file_path)
        files_content.append((os.path.basename(file_path), content))

    # 创建提示词 - 准备提供给LLM的输入
    print("正在准备提示词...")
    prompt = create_prompt(files_content, args.instruction)

    # 调用LLM进行融合重写 - 调用API并显示处理状态
    print("正在调用LLM进行融合重写...")
    result = call_llm(prompt)

    # 处理API返回结果 - 如果成功则保存文档，否则显示错误信息
    if result:
        # 保存结果
        save_to_html(result, args.output)
    else:
        print("融合失败，未能获取有效内容")

if __name__ == "__main__":
    # 程序入口点 - 调用主函数
    main()

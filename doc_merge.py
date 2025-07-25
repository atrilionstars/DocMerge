import argparse
import json
import os
from typing import List, Tuple

import markdown
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from tqdm import tqdm

BASE_URL = "http://192.168.1.13:8000/v1"
API_KEY = "2JysWWdHfyKvp2AsGYznw7pwPfkwDehtPZHEtj26GIA"
MODEL_NAME = "qwen-max"

# 请求头
headers = {
    "Content-Type": "application/json",
    "Authorization": f"Bearer {API_KEY}"
}

def read_table(table) -> str:
    """读取表格内容并转换为文本格式"""
    table_text = []
    # 记录表头和内容分隔符
    table_text.append("=== 表格开始 ===")

    # 遍历表格的每一行
    for row_idx, row in enumerate(table.rows):
        row_data = []
        # 遍历行中的每个单元格
        for cell in row.cells:
            # 提取单元格文本，处理可能的换行
            cell_text = "\n".join([p.text for p in cell.paragraphs]).strip()
            row_data.append(cell_text)

        # 用|分隔单元格内容，便于识别
        table_text.append("|".join(row_data))

        # 在表头下方添加分隔线（如果有表头）
        if row_idx == 0 and len(table.rows) > 1:
            table_text.append("|".join(["---"] * len(row_data)))

    table_text.append("=== 表格结束 ===")
    return "\n".join(table_text)

def read_docx(file_path: str) -> str:
    """读取docx文档内容，包括段落和表格"""
    # 创建Document对象读取Word文档
    doc = Document(file_path)
    full_text = []

    # 遍历文档中的每个元素，区分段落和表格
    for element in doc.element.body:
        # 检查元素是否为段落
        if element.tag.endswith('p'):
            # 找到对应的段落对象
            para = next((p for p in doc.paragraphs if p._element == element), None)
            if para and para.text.strip():
                full_text.append(para.text)
        # 检查元素是否为表格
        elif element.tag.endswith('tbl'):
            # 找到对应的表格对象
            table = next((t for t in doc.tables if t._element == element), None)
            if table:
                # 读取表格内容并添加到文本中
                table_content = read_table(table)
                full_text.append(table_content)

    # 将所有内容用换行符连接成一个字符串并返回
    return '\n'.join(full_text)

def create_prompt(files_content: List[Tuple[str, str]], instruction: str = "") -> str:
    """创建提示词"""
    # 初始化提示词，介绍后续是多个文档的内容
    prompt = "以下是多个文档的内容:\n\n"

    # 遍历每个文档，将文档名和内容添加到提示词中
    for i, (filename, content) in enumerate(files_content):
        prompt += f"文档 {i+1} ({filename}):\n{content}\n\n"

    # 如果提供了额外说明，将其添加到提示词中
    if instruction:
        prompt += f"额外说明: {instruction}\n\n"

    # 要求模型根据以下格式进行融合重写
    prompt += "请将我提供的多个Word文档内容进行整合重写，要求如下：1.提取各文档的核心信息。2.保留所有文档中的独特数据和观点。3.按逻辑重新组织章节结构"
    return prompt

def call_llm(prompt: str) -> str:
    """调用LLM API生成内容"""

    url = f"{BASE_URL}/chat/completions"
    payload = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system", "content":
                "你是一个专业的文职秘书，负责公司的撰写工作，勿用口语化表述。"},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3
    }

    try:
        print(f"发送请求到API，提示词长度: {len(prompt)} 字符")
        response = requests.post(url, headers=headers, data=json.dumps(payload))
        response.raise_for_status()
        result = response.json()["choices"][0]["message"]["content"].strip()

        # 调试输出
        print(f"API返回内容长度: {len(result)} 字符")
        if len(result) < 10:  # 如果内容非常短，可能有问题
            print(f"警告: API返回的内容异常短: {result[:50]}...")

        return result
    except Exception as e:
        print(f"API调用错误: {e}")
        return None

def md_content_to_word(md_content, output_filename="output.docx"):
    """
    将Markdown字符串内容转换为Word文档并保存到当前工作路径，支持更多格式和样式

    参数:
        md_content: Markdown格式的字符串内容
        output_filename: 输出Word文档文件名，默认为"output.docx"
    """
    # 创建Word文档
    doc = Document()

    # 设置页面属性
    section = doc.sections[0]
    # 设置页边距
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    # 设置纸张大小为A4
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    # 正文样式
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = '宋体'
    normal_font.size = Pt(11)
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 设置段落格式
    normal_style.paragraph_format.line_spacing = 1.5  # 行间距1.5倍
    normal_style.paragraph_format.space_after = Pt(0)  # 段后间距
    normal_style.paragraph_format.first_line_indent = Pt(22)  # 首行缩进2字符

    # 标题样式 - 根据Markdown标题级别设置不同大小
    title_sizes = [Pt(20), Pt(17), Pt(15), Pt(13), Pt(11), Pt(11)]  # 对应h1到h6
    title_spacing = [Pt(12), Pt(10), Pt(8), Pt(6), Pt(5), Pt(5)]  # 标题段后间距

    for level in range(1, 7):
        heading_style = doc.styles[f'Heading {level}']
        heading_font = heading_style.font
        heading_font.name = '宋体'
        heading_font.size = title_sizes[level-1]
        heading_font.bold = False # 标题加粗
        heading_font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        # 设置中文字体
        heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 标题左对齐
        heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 设置标题间距
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = title_spacing[level-1]
        heading_style.paragraph_format.line_spacing = 1.2

    # 将Markdown转换为HTML，然后使用BeautifulSoup解析
    html = markdown.markdown(
        md_content,
        extensions=[
            'extra',  # 支持表格、代码块等扩展语法
            'codehilite',  # 代码高亮
            'tables'  # 表格支持
        ]
    )
    soup = BeautifulSoup(html, 'html.parser')

    # 处理解析后的HTML内容
    for element in soup.contents:
        if element.name is None:
            continue

        # 处理标题
        if element.name.startswith('h'):
            try:
                level = int(element.name[1:])
                if 1 <= level <= 6:
                    text = element.get_text()
                    doc.add_heading(text, level=level)
            except (ValueError, IndexError):
                pass
            continue

        # 处理段落
        if element.name == 'p':
            para = doc.add_paragraph()
            para.style = 'Normal'
            # 处理段落内的各种元素
            for child in element.contents:
                process_inline_element(child, para)
            continue

        # 处理无序列表
        if element.name == 'ul':
            process_list(element, doc, is_ordered=False)
            continue

        # 处理有序列表
        if element.name == 'ol':
            process_list(element, doc, is_ordered=True)
            continue

        # 处理表格
        if element.name == 'table':
            process_table(element, doc)
            continue


    # 获取当前工作路径并保存Word文档
    output_path = os.path.join(os.getcwd(), output_filename)
    doc.save(output_path)
    print(f"转换完成，文件已保存至: {output_path}")
    return output_path

def process_inline_element(element, paragraph):
    """处理行内元素（如粗体、斜体、链接等）"""
    if isinstance(element, str):
        if element.strip():
            paragraph.add_run(element.strip())
        return

    if element.name is None:
        return

    # 处理粗体
    if element.name == 'strong' or element.name == 'b':
        run = paragraph.add_run(element.get_text())
        run.bold = True
        return

    # 处理斜体
    if element.name == 'em' or element.name == 'i':
        run = paragraph.add_run(element.get_text())
        run.italic = True
        return

    # 处理链接
    if element.name == 'a':
        url = element.get('href', '')
        text = element.get_text() or url
        run = paragraph.add_run(text)
        run.underline = True
        run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
        # 注意：python-docx不直接支持超链接，这里仅设置样式
        return

    # 处理代码（行内）
    if element.name == 'code':
        run = paragraph.add_run(element.get_text())
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        return

    # 处理其他行内元素
    for child in element.contents:
        process_inline_element(child, paragraph)

def process_list(list_element, doc, is_ordered=False):
    """处理列表（有序和无序）"""
    list_style = 'List Number' if is_ordered else 'List Bullet'
    current_para = None

    for item in list_element.find_all('li', recursive=False):
        # 创建新列表项
        para = doc.add_paragraph(style=list_style)

        # 处理列表项内容
        for child in item.contents:
            # 处理嵌套列表
            if child.name in ['ul', 'ol']:
                nested_is_ordered = child.name == 'ol'
                process_list(child, doc, nested_is_ordered)
            else:
                process_inline_element(child, para)

        current_para = para

def process_table(table_element, doc):
    """处理表格"""
    # 获取表格行数和列数
    rows = table_element.find_all('tr')
    if not rows:
        return

    # 创建Word表格
    table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['th', 'td'])))
    table.style = 'Table Grid'  # 使用网格样式

    # 填充表格内容
    for row_idx, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for col_idx, cell in enumerate(cells):
            if col_idx < len(table.columns):
                table_cell = table.cell(row_idx, col_idx)
                # 表头单元格加粗
                if cell.name == 'th':
                    para = table_cell.add_paragraph()
                    run = para.add_run(cell.get_text())
                    run.bold = True
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    para = table_cell.add_paragraph()
                    process_inline_element(cell, para)

    # 添加表格后的空行
    doc.add_paragraph()

def main():
    # 创建命令行参数解析器，设置工具描述
    parser = argparse.ArgumentParser(description="多文档融合重写工具")
    # 添加输入文件参数 - 允许多个docx文件路径
    parser.add_argument("-i", "--input", nargs="+", help="输入的docx文件路径")
    # 添加输出文件参数 - 指定融合后的文档保存路径
    parser.add_argument("-o", "--output", help="输出文件的路径")
    # 添加额外说明参数 - 可以提供给模型的额外指令
    parser.add_argument("-t", "--instruction", help="额外说明或指令")
    # 解析命令行参数
    args = parser.parse_args()
    # 处理默认输入文件 - 如果未指定输入文件，使用当前目录下的所有docx文件
    if not args.input:
        print("未指定输入文件，将使用当前目录下的所有docx文件")
        current_dir = os.getcwd()
        args.input = [os.path.join(current_dir, f) for f in os.listdir(current_dir) if f.lower().endswith('.docx')]

        # 检查是否找到docx文件
        if not args.input:
            print("错误: 当前目录下没有找到docx文件")
            return

    # 处理默认输出文件 - 如果未指定输出文件，使用默认名称
    if not args.output:
        print("未指定输出文件，将使用默认名称")
        current_dir = os.getcwd()
        args.output = os.path.join(current_dir, "融合文档.docx")

    # 检查输入文件是否存在 - 收集不存在的文件
    invalid_files = [f for f in args.input if not os.path.exists(f)]
    if invalid_files:
        print(f"错误: 以下文件不存在: {', '.join(invalid_files)}")
        return

    # 检查输出文件目录是否存在 - 如果不存在则创建
    output_dir = os.path.dirname(args.output)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 读取所有文档内容 - 使用进度条显示读取进度
    print("正在读取文档...")
    files_content = []
    for file_path in tqdm(args.input):
        content = read_docx(file_path)
        files_content.append((os.path.basename(file_path), content))

    # 创建提示词 - 准备提供给LLM的输入
    print("正在准备提示词...")
    prompt = create_prompt(files_content, args.instruction)

    # 调用LLM进行融合重写 - 调用API并显示处理状态
    print("正在调用LLM进行融合重写...")
    result = call_llm(prompt)

    # 处理API返回结果 - 如果成功则保存文档，否则显示错误信息
    if result:
        # 保存结果
        md_content_to_word(result, args.output)
    else:
        print("融合失败，未能获取有效内容")

if __name__ == "__main__":
    main()

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import argparse

def convert_text_to_docx(input_text, output_file):
    """
    将包含标题标记的文本转换为带编号的docx文档

    参数:
        input_text: 输入的文本内容
        output_file: 输出的docx文件路径
    """
    # 创建一个新的Word文档
    doc = Document()

    # 定义标题样式的字典，键是标题级别，值是字体大小
    title_styles = {
        1: 24,   # 一级标题
        2: 20,   # 二级标题
        3: 16,   # 三级标题
        4: 14,   # 四级标题
        5: 12    # 五级标题
    }

    # 用于跟踪标题编号
    title_numbers = [0, 0, 0, 0, 0]  # 分别对应1-5级标题的编号

    # 按行分割文本
    lines = input_text.split('\n')

    for line in lines:
        # 去除行首尾的空白字符
        line = line.strip()

        if not line:  # 跳过空行
            continue

        # 检查是否是标题行（以#开头）
        title_match = re.match(r'^(#{1,5})\s+(.*)$', line)
        if title_match:
            # 获取标题级别和标题内容
            title_level = len(title_match.group(1))
            title_text = title_match.group(2)

            # 更新标题编号
            title_numbers[title_level - 1] += 1
            # 重置更低级别的标题编号
            for i in range(title_level, 5):
                title_numbers[i] = 0

            # 生成标题编号字符串
            number_str = '.'.join(str(num) for num in title_numbers[:title_level] if num > 0) + ' '

            # 添加带编号的标题
            if title_level in title_styles:
                para = doc.add_paragraph(number_str + title_text)
                # 设置标题样式
                run = para.runs[0]
                run.font.size = Pt(title_styles[title_level])
                run.font.bold = True
                # 标题居中对齐
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # 添加正文内容
            para = doc.add_paragraph(line)
            # 设置正文字体和大小
            run = para.runs[0]
            run.font.size = Pt(12)

    # 保存文档
    doc.save(output_file)
    # 保存一个文本版本用于对比
    text_output_file = output_file.replace('.docx', '_debug.txt')
    with open(text_output_file, 'w', encoding='utf-8') as f:
        f.write(input_text)

    print(f"成功将内容转换为docx文档：{output_file}")

def main():
    # 设置命令行参数
    parser = argparse.ArgumentParser(description='将大模型输出的文本转换为带标题编号的docx格式')
    parser.add_argument('-i', '--input', required=True, help='输入的文本文件路径')
    parser.add_argument('-o', '--output', required=True, help='输出的docx文件路径')

    args = parser.parse_args()

    # 读取输入文件
    try:
        with open(args.input, 'r', encoding='utf-8') as f:
            input_text = f.read()

        # 转换并保存为docx
        convert_text_to_docx(input_text, args.output)

    except Exception as e:
        print(f"转换过程中出现错误：{str(e)}")

if __name__ == "__main__":
    main()
